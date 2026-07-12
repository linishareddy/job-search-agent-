"""Render structured resume sections into a Classic DOCX template."""
import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from schemas.resume_sections import ParsedResumeSections, TailoredResumeSections


def sections_to_plain_text(sections: ParsedResumeSections | TailoredResumeSections | dict[str, Any]) -> str:
    """Flatten structured sections into plain text for backward-compatible display."""
    if isinstance(sections, dict):
        sections = ParsedResumeSections.model_validate(sections)

    lines: list[str] = []
    c = sections.contact
    contact_bits = [c.name, c.email, c.phone, c.location, c.linkedin, c.github, c.website]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        lines.append(contact_line)
        lines.append("")

    if sections.summary:
        lines.append("SUMMARY")
        lines.append(sections.summary)
        lines.append("")

    if sections.skills:
        lines.append("SKILLS")
        lines.append(", ".join(sections.skills))
        lines.append("")

    if sections.experience:
        lines.append("EXPERIENCE")
        for exp in sections.experience:
            header = " — ".join(
                p for p in [exp.title, exp.company, exp.location] if p
            )
            dates = " – ".join(p for p in [exp.start_date, exp.end_date] if p)
            if header:
                lines.append(header)
            if dates:
                lines.append(dates)
            for bullet in exp.bullets:
                lines.append(f"• {bullet}")
            lines.append("")

    if sections.education:
        lines.append("EDUCATION")
        for edu in sections.education:
            header = " — ".join(p for p in [edu.degree, edu.institution, edu.location] if p)
            if header:
                lines.append(header)
            if edu.graduation_date:
                lines.append(edu.graduation_date)
            for detail in edu.details:
                lines.append(f"• {detail}")
            lines.append("")

    if sections.certifications:
        lines.append("CERTIFICATIONS")
        for cert in sections.certifications:
            lines.append(f"• {cert}")

    return "\n".join(lines).strip()


def render_classic_docx(sections: TailoredResumeSections | dict[str, Any]) -> bytes:
    """Build a single-column Classic resume DOCX from structured sections."""
    if isinstance(sections, dict):
        sections = TailoredResumeSections.model_validate(sections)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    c = sections.contact
    if c.name:
        name_para = doc.add_paragraph(c.name)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.runs[0].bold = True
        name_para.runs[0].font.size = Pt(16)

    contact_bits = [c.email, c.phone, c.location, c.linkedin, c.github, c.website]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        contact_para = doc.add_paragraph(contact_line)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)

    def add_section_heading(title: str) -> None:
        para = doc.add_paragraph()
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(12)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)

    if sections.summary:
        add_section_heading("Summary")
        doc.add_paragraph(sections.summary)

    if sections.skills:
        add_section_heading("Skills")
        doc.add_paragraph(", ".join(sections.skills))

    if sections.experience:
        add_section_heading("Experience")
        for exp in sections.experience:
            title_parts = [exp.title, exp.company]
            title_line = " — ".join(p for p in title_parts if p)
            if title_line:
                p = doc.add_paragraph()
                run = p.add_run(title_line)
                run.bold = True
            dates = " – ".join(p for p in [exp.start_date, exp.end_date] if p)
            if dates:
                date_para = doc.add_paragraph(dates)
                date_para.runs[0].italic = True
                date_para.runs[0].font.size = Pt(10)
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if sections.education:
        add_section_heading("Education")
        for edu in sections.education:
            header = " — ".join(p for p in [edu.degree, edu.institution] if p)
            if header:
                p = doc.add_paragraph()
                p.add_run(header).bold = True
            if edu.graduation_date:
                doc.add_paragraph(edu.graduation_date)
            for detail in edu.details:
                doc.add_paragraph(detail, style="List Bullet")

    if sections.certifications:
        add_section_heading("Certifications")
        for cert in sections.certifications:
            doc.add_paragraph(cert, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
